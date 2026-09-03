"""Contract tests for the deterministic A-to-Z analysis pipeline."""
import io
import json
import unittest

import pandas as pd

from utils.analysis_pipeline import (
    DataLoadError,
    analyze_dataframe,
    analyze_document_text,
    apply_cleaning_plan,
    export_csv_safe,
    load_source,
    profile_dataframe,
    recommend_cleaning,
    render_markdown_report,
)


class AnalysisPipelineTests(unittest.TestCase):

 def test_csv_ingestion_detects_semicolon_and_cp1252(self):
    raw = "Customer;Amount\nJos\xe9;1,200\nAda;900\n".encode("cp1252")
    result = load_source(raw, "sales.csv")
    self.assertEqual(result.metadata["delimiter"], ";")
    self.assertEqual(result.primary.loc[0, "Customer"], "José")
    self.assertEqual(result.primary.shape, (2, 2))


 def test_json_ingestion_returns_each_record_collection(self):
    payload = {"sales": [{"id": 1, "value": 3}], "people": [{"name": "Ada"}]}
    result = load_source(json.dumps(payload).encode(), "bundle.json")
    self.assertEqual(set(result.datasets), {"sales", "people"})
    self.assertEqual(result.datasets["sales"].loc[0, "value"], 3)


 def test_excel_ingestion_preserves_multiple_sheets(self):
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine="openpyxl") as writer:
        pd.DataFrame({"a": [1]}).to_excel(writer, sheet_name="First sheet", index=False)
        pd.DataFrame({"b": [2]}).to_excel(writer, sheet_name="Second", index=False)
    result = load_source(output.getvalue(), "book.xlsx")
    self.assertEqual(set(result.datasets), {"First_sheet", "Second"})

 def test_pdf_ingestion_extracts_selectable_text(self):
    import fitz

    document = fitz.open()
    page = document.new_page()
    page.insert_text((72, 72), "Revenue increased 18 percent in 2026.")
    raw = document.tobytes()
    document.close()
    result = load_source(raw, "report.pdf")
    self.assertIn("Revenue increased", result.document_text)
    self.assertEqual(result.metadata["pages"], 1)

 def test_docx_ingestion_extracts_text_and_tables(self):
    from docx import Document

    output = io.BytesIO()
    document = Document()
    document.add_paragraph("Quarterly operating review")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Region"
    table.cell(0, 1).text = "Revenue"
    table.cell(1, 0).text = "West"
    table.cell(1, 1).text = "1200"
    document.save(output)
    result = load_source(output.getvalue(), "review.docx")
    self.assertIn("Quarterly operating review", result.document_text)
    self.assertEqual(result.datasets["table_1"].loc[0, "Region"], "West")


 def test_profile_finds_missing_duplicates_types_and_outliers(self):
    df = pd.DataFrame({
        "id": [1, 2, 2, 4, 5],
        "amount": [10.0, 11.0, 11.0, None, 10_000.0],
        "date": ["2026-01-01", "2026-01-02", "2026-01-02", "2026-01-04", "2026-01-05"],
    })
    profile = profile_dataframe(df)
    by_name = {item["column"]: item for item in profile["column_profiles"]}
    self.assertEqual(profile["duplicate_rows"], 1)
    self.assertEqual(profile["missing_cells"], 1)
    self.assertEqual(by_name["date"]["semantic_type"], "datetime_text")
    self.assertEqual(by_name["amount"]["outliers_iqr"], 1)


 def test_cleaning_is_allowlisted_conservative_and_audited(self):
    df = pd.DataFrame({
        " Customer Name ": [" Ada ", " Ada ", "Grace", "-"],
        "Amount USD": ["1,200", "1,200", "900", "null"],
    })
    profile = profile_dataframe(df)
    actions = recommend_cleaning(df, profile)
    selected = [item for item in actions if item["action"] in {"normalize_columns", "normalize_missing", "trim_text"}]
    selected.append({"action": "convert_numeric", "column": "amount_usd"})
    cleaned, audit = apply_cleaning_plan(df, selected)
    self.assertEqual(list(cleaned.columns), ["customer_name", "amount_usd"])
    self.assertEqual(cleaned.loc[0, "customer_name"], "Ada")
    self.assertEqual(cleaned.loc[0, "amount_usd"], 1200)
    self.assertTrue(pd.isna(cleaned.loc[3, "amount_usd"]))
    self.assertEqual([entry["step"] for entry in audit], [1, 2, 3, 4])
    with self.assertRaises(ValueError):
        apply_cleaning_plan(df, [{"action": "run_python"}])


 def test_analysis_reports_evidence_and_causal_caveat(self):
    df = pd.DataFrame({"x": range(20), "y": [value * 2 for value in range(20)], "group": ["A"] * 12 + ["B"] * 8})
    analysis = analyze_dataframe(df, "Understand the strongest patterns")
    self.assertEqual(analysis["correlations"][0]["correlation"], 1.0)
    self.assertIn("does not establish causation", analysis["findings"][0]["caveat"])
    report = render_markdown_report("test.csv", "data", analysis, [])
    self.assertIn("DataPrism Analysis Report", report)
    self.assertIn("Pearson r=1.00", report)
    self.assertIn("Limitations", report)


 def test_document_analysis_is_extractive(self):
    text = (
        "Revenue increased by 18% in 2026 after three quarters of sustained growth. "
        "Revenue growth was strongest in the enterprise customer segment, reaching $1,250,000."
    )
    result = analyze_document_text(text)
    self.assertEqual(result["number_count"], 3)
    self.assertEqual(result["top_terms"]["revenue"], 2)
    self.assertTrue(all(sentence in text for sentence in result["key_sentences"]))


 def test_unsupported_and_empty_files_fail_clearly(self):
    with self.assertRaisesRegex(DataLoadError, "empty"):
        load_source(b"", "empty.csv")
    with self.assertRaisesRegex(DataLoadError, "Unsupported"):
        load_source(b"hello", "data.exe")

 def test_csv_export_neutralizes_formula_injection(self):
    exported = export_csv_safe(pd.DataFrame({"name": ["Ada"], "note": ["=HYPERLINK(\"bad\")"], "amount": [-3]})).decode("utf-8-sig")
    self.assertIn("'=HYPERLINK", exported)
    self.assertIn(",-3", exported)
