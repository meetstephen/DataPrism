"""Deterministic, auditable analysis pipeline used by DataPrism's workbench.

The functions in this module deliberately do not depend on Streamlit.  They can
therefore be tested in isolation and reused by every UI page.  AI may explain
the results later, but ingestion, cleaning, metrics and evidence are computed
locally from the data.
"""
from __future__ import annotations

import csv
import io
import json
import math
import re
from dataclasses import dataclass, field
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, BinaryIO, Iterable

import numpy as np
import pandas as pd


MAX_UPLOAD_BYTES = 200 * 1024 * 1024
SUPPORTED_EXTENSIONS = {
    ".csv", ".tsv", ".txt", ".xlsx", ".xls", ".json", ".jsonl",
    ".ndjson", ".parquet", ".pdf", ".docx", ".md",
}
MISSING_TOKENS = {"", "-", "--", "n/a", "na", "null", "none", "nil", "nan", "missing"}


class DataLoadError(ValueError):
    """A user-correctable ingestion error with a safe display message."""


@dataclass
class LoadResult:
    source_name: str
    source_format: str
    datasets: dict[str, pd.DataFrame] = field(default_factory=dict)
    document_text: str = ""
    warnings: list[str] = field(default_factory=list)
    metadata: dict[str, Any] = field(default_factory=dict)

    @property
    def primary(self) -> pd.DataFrame | None:
        return next(iter(self.datasets.values()), None)


def _read_bytes(source: bytes | bytearray | BinaryIO) -> bytes:
    if isinstance(source, (bytes, bytearray)):
        raw = bytes(source)
    else:
        try:
            source.seek(0)
        except (AttributeError, OSError):
            pass
        raw = source.read()
        if isinstance(raw, str):
            raw = raw.encode("utf-8")
    if not raw:
        raise DataLoadError("The file is empty.")
    if len(raw) > MAX_UPLOAD_BYTES:
        raise DataLoadError(
            f"The file is {len(raw) / 1024 / 1024:.1f} MB; the limit is "
            f"{MAX_UPLOAD_BYTES / 1024 / 1024:.0f} MB."
        )
    return raw


def _decode_text(raw: bytes) -> tuple[str, str]:
    candidates = ["utf-8-sig", "utf-8", "cp1252", "latin-1"]
    try:
        import chardet
        guess = chardet.detect(raw[:250_000]).get("encoding")
        if guess:
            candidates.insert(0, guess)
    except ImportError:
        pass
    seen: set[str] = set()
    for encoding in candidates:
        if encoding.lower() in seen:
            continue
        seen.add(encoding.lower())
        try:
            return raw.decode(encoding), encoding
        except (UnicodeDecodeError, LookupError):
            continue
    raise DataLoadError("The text encoding could not be detected. Re-save the file as UTF-8.")


def _clean_dataset_name(value: str, fallback: str = "dataset") -> str:
    value = re.sub(r"[^A-Za-z0-9_-]+", "_", str(value)).strip("_")
    return value[:80] or fallback


def _dedupe_columns(columns: Iterable[Any]) -> list[str]:
    seen: dict[str, int] = {}
    result: list[str] = []
    for index, value in enumerate(columns, 1):
        base = str(value).strip() if value is not None else ""
        if not base or base.lower().startswith("unnamed:"):
            base = f"column_{index}"
        count = seen.get(base, 0)
        seen[base] = count + 1
        result.append(base if count == 0 else f"{base}_{count + 1}")
    return result


def _finalize_frame(df: pd.DataFrame) -> pd.DataFrame:
    if df is None or df.empty or len(df.columns) == 0:
        raise DataLoadError("The file was read, but it contains no tabular data.")
    result = df.copy()
    result.columns = _dedupe_columns(result.columns)
    result = result.dropna(axis=0, how="all").dropna(axis=1, how="all")
    if result.empty or len(result.columns) == 0:
        raise DataLoadError("The file contains only blank rows or columns.")
    return result.reset_index(drop=True)


def _read_delimited(raw: bytes, forced_separator: str | None = None) -> tuple[pd.DataFrame, dict[str, Any]]:
    text, encoding = _decode_text(raw)
    sample = text[:64_000]
    separator = forced_separator
    if separator is None:
        try:
            separator = csv.Sniffer().sniff(sample, delimiters=",\t;|").delimiter
        except csv.Error:
            separator = ","
    try:
        df = pd.read_csv(io.StringIO(text), sep=separator, engine="python")
    except Exception as exc:
        raise DataLoadError(f"The delimited file could not be parsed: {exc}") from exc
    return _finalize_frame(df), {"encoding": encoding, "delimiter": separator}


def _read_json(raw: bytes, extension: str) -> dict[str, pd.DataFrame]:
    text, _ = _decode_text(raw)
    try:
        if extension in {".jsonl", ".ndjson"}:
            return {"records": _finalize_frame(pd.read_json(io.StringIO(text), lines=True))}
        payload = json.loads(text)
    except (ValueError, json.JSONDecodeError) as exc:
        raise DataLoadError(f"The JSON is invalid: {exc}") from exc

    datasets: dict[str, pd.DataFrame] = {}
    if isinstance(payload, list):
        datasets["records"] = _finalize_frame(pd.json_normalize(payload))
    elif isinstance(payload, dict):
        for key, value in payload.items():
            if isinstance(value, list) and value and all(isinstance(row, dict) for row in value):
                datasets[_clean_dataset_name(key)] = _finalize_frame(pd.json_normalize(value))
        if not datasets:
            datasets["record"] = _finalize_frame(pd.json_normalize(payload))
    else:
        raise DataLoadError("JSON must contain an object or an array of records.")
    return datasets


def _table_from_rows(rows: list[list[Any]]) -> pd.DataFrame | None:
    cleaned = [list(row) for row in rows if row and any(cell not in (None, "") for cell in row)]
    if len(cleaned) < 2:
        return None
    width = max(len(row) for row in cleaned)
    cleaned = [row + [None] * (width - len(row)) for row in cleaned]
    first = ["" if value is None else str(value).strip() for value in cleaned[0]]
    nonempty = sum(bool(value) for value in first)
    unique = len({value.lower() for value in first if value}) == nonempty
    text_cells = sum(bool(value) and not re.fullmatch(r"[+-]?[\d,.%]+", value) for value in first)
    header_like = nonempty >= max(1, math.ceil(width / 2)) and unique and text_cells >= math.ceil(nonempty / 2)
    if header_like:
        df = pd.DataFrame(cleaned[1:], columns=_dedupe_columns(first))
    else:
        df = pd.DataFrame(cleaned, columns=[f"column_{i}" for i in range(1, width + 1)])
    try:
        return _finalize_frame(df)
    except DataLoadError:
        return None


def _read_pdf(raw: bytes, result: LoadResult) -> None:
    try:
        import fitz
    except ImportError as exc:
        raise DataLoadError("PDF support requires PyMuPDF.") from exc
    try:
        document = fitz.open(stream=raw, filetype="pdf")
    except Exception as exc:
        raise DataLoadError(f"The PDF is encrypted, damaged, or unreadable: {exc}") from exc
    pages: list[str] = []
    try:
        for page_index, page in enumerate(document, 1):
            text = page.get_text("text").strip()
            if text:
                pages.append(f"--- Page {page_index} ---\n{text}")
            try:
                finder = page.find_tables()
                for table_index, table in enumerate(finder.tables, 1):
                    frame = _table_from_rows(table.extract())
                    if frame is not None:
                        result.datasets[f"page_{page_index}_table_{table_index}"] = frame
            except Exception as exc:
                result.warnings.append(f"Table detection failed on page {page_index}: {exc}")
        result.document_text = "\n\n".join(pages)
        result.metadata["pages"] = len(document)
    finally:
        document.close()
    if not result.document_text:
        result.warnings.append(
            "No selectable text was found. This may be a scanned PDF; OCR is not currently performed."
        )
    if not result.datasets:
        result.warnings.append(
            "No reliable tables were detected. The PDF is available for document analysis only."
        )


def _read_docx(raw: bytes, result: LoadResult) -> None:
    try:
        from docx import Document
    except ImportError as exc:
        raise DataLoadError("Word support requires python-docx.") from exc
    try:
        document = Document(io.BytesIO(raw))
    except Exception as exc:
        raise DataLoadError(f"The Word document could not be read: {exc}") from exc
    result.document_text = "\n\n".join(p.text for p in document.paragraphs if p.text.strip())
    for table_index, table in enumerate(document.tables, 1):
        rows = [[cell.text for cell in row.cells] for row in table.rows]
        frame = _table_from_rows(rows)
        if frame is not None:
            result.datasets[f"table_{table_index}"] = frame
    result.metadata["paragraphs"] = len(document.paragraphs)
    result.metadata["tables"] = len(document.tables)


def load_source(source: bytes | bytearray | BinaryIO, filename: str) -> LoadResult:
    """Load a supported file into one or more datasets plus document context."""
    raw = _read_bytes(source)
    extension = Path(filename).suffix.lower()
    if extension not in SUPPORTED_EXTENSIONS:
        supported = ", ".join(sorted(SUPPORTED_EXTENSIONS))
        raise DataLoadError(f"Unsupported file type '{extension or 'unknown'}'. Supported: {supported}")
    result = LoadResult(
        source_name=Path(filename).name,
        source_format=extension.lstrip("."),
        metadata={"bytes": len(raw), "loaded_at": datetime.now(timezone.utc).isoformat()},
    )
    if extension in {".csv", ".tsv"}:
        frame, metadata = _read_delimited(raw, "\t" if extension == ".tsv" else None)
        result.datasets["data"] = frame
        result.metadata.update(metadata)
    elif extension in {".txt", ".md"}:
        text, encoding = _decode_text(raw)
        result.document_text = text
        result.metadata["encoding"] = encoding
        try:
            frame, metadata = _read_delimited(raw)
            if len(frame.columns) > 1:
                result.datasets["data"] = frame
                result.metadata.update(metadata)
        except DataLoadError:
            pass
    elif extension in {".xlsx", ".xls"}:
        try:
            workbook = pd.ExcelFile(io.BytesIO(raw))
            for sheet in workbook.sheet_names:
                frame = pd.read_excel(workbook, sheet_name=sheet)
                try:
                    result.datasets[_clean_dataset_name(sheet, "sheet")] = _finalize_frame(frame)
                except DataLoadError:
                    result.warnings.append(f"Skipped empty sheet: {sheet}")
            result.metadata["sheets"] = workbook.sheet_names
        except Exception as exc:
            raise DataLoadError(f"The Excel workbook could not be read: {exc}") from exc
    elif extension in {".json", ".jsonl", ".ndjson"}:
        result.datasets.update(_read_json(raw, extension))
    elif extension == ".parquet":
        try:
            result.datasets["data"] = _finalize_frame(pd.read_parquet(io.BytesIO(raw)))
        except Exception as exc:
            raise DataLoadError(f"The Parquet file could not be read: {exc}") from exc
    elif extension == ".pdf":
        _read_pdf(raw, result)
    elif extension == ".docx":
        _read_docx(raw, result)
    if not result.datasets and not result.document_text:
        raise DataLoadError("No usable data or text could be extracted from this file.")
    result.metadata["dataset_count"] = len(result.datasets)
    return result


def _safe_number(value: Any) -> float | int | None:
    if value is None or pd.isna(value) or not np.isfinite(value):
        return None
    return int(value) if float(value).is_integer() else float(value)


def infer_semantic_type(series: pd.Series) -> tuple[str, float]:
    """Infer a useful analytical role and return (role, confidence)."""
    non_null = series.dropna()
    if non_null.empty:
        return "empty", 1.0
    name = str(series.name).lower()
    if pd.api.types.is_bool_dtype(series):
        return "boolean", 1.0
    if pd.api.types.is_datetime64_any_dtype(series):
        return "datetime", 1.0
    if pd.api.types.is_numeric_dtype(series):
        uniqueness = non_null.nunique() / len(non_null)
        if uniqueness > 0.98 and re.search(r"(^id$|_id$|^id_|code$|number$)", name):
            return "identifier", 0.95
        return "numeric", 1.0
    sample = non_null.astype(str).str.strip().head(2_000)
    lowered = sample.str.lower()
    if set(lowered.unique()).issubset({"true", "false", "yes", "no", "y", "n", "0", "1"}):
        return "boolean", 0.9
    numeric_ratio = pd.to_numeric(sample.str.replace(",", "", regex=False), errors="coerce").notna().mean()
    if numeric_ratio >= 0.95:
        return "numeric_text", float(numeric_ratio)
    date_ratio = pd.to_datetime(sample, errors="coerce", format="mixed").notna().mean()
    if date_ratio >= 0.9:
        return "datetime_text", float(date_ratio)
    unique = sample.nunique()
    if unique <= 50 or unique / len(sample) <= 0.1:
        return "categorical", 0.85
    if unique == len(sample) and re.search(r"(^id$|_id$|^id_|code$|email$)", name):
        return "identifier", 0.85
    return "text", 0.8


def profile_dataframe(df: pd.DataFrame) -> dict[str, Any]:
    """Produce a JSON-safe quality and schema profile with explicit evidence."""
    if df is None or df.empty:
        raise ValueError("A non-empty DataFrame is required.")
    rows, columns = df.shape
    total_cells = rows * columns
    missing_cells = int(df.isna().sum().sum())
    duplicate_rows = int(df.duplicated().sum())
    column_profiles: list[dict[str, Any]] = []
    issues: list[dict[str, Any]] = []
    for column in df.columns:
        series = df[column]
        non_null = series.dropna()
        role, confidence = infer_semantic_type(series)
        missing = int(series.isna().sum())
        unique = int(series.nunique(dropna=True))
        item: dict[str, Any] = {
            "column": str(column), "dtype": str(series.dtype), "semantic_type": role,
            "type_confidence": round(confidence, 2), "non_null": rows - missing,
            "missing": missing, "missing_pct": round(missing / rows * 100, 2),
            "unique": unique, "unique_pct": round(unique / max(rows - missing, 1) * 100, 2),
        }
        if pd.api.types.is_numeric_dtype(series) and not non_null.empty:
            clean = pd.to_numeric(non_null, errors="coerce").dropna()
            q1, q3 = clean.quantile([0.25, 0.75])
            iqr = q3 - q1
            outliers = int(((clean < q1 - 1.5 * iqr) | (clean > q3 + 1.5 * iqr)).sum()) if iqr > 0 else 0
            item.update({
                "mean": _safe_number(clean.mean()), "median": _safe_number(clean.median()),
                "std": _safe_number(clean.std()), "min": _safe_number(clean.min()),
                "max": _safe_number(clean.max()), "outliers_iqr": outliers,
            })
        else:
            item["top_values"] = {str(k): int(v) for k, v in non_null.astype(str).value_counts().head(5).items()}
        column_profiles.append(item)
        if missing / rows >= 0.4:
            issues.append({"severity": "high", "code": "high_missing", "column": str(column), "evidence": f"{missing / rows:.1%} missing"})
        elif missing:
            issues.append({"severity": "medium", "code": "missing_values", "column": str(column), "evidence": f"{missing} missing values"})
        if unique <= 1:
            issues.append({"severity": "medium", "code": "constant_column", "column": str(column), "evidence": f"{unique} distinct value"})
        if role in {"numeric_text", "datetime_text"}:
            issues.append({"severity": "medium", "code": "convertible_type", "column": str(column), "evidence": f"{role} ({confidence:.0%} parseable)"})
    if duplicate_rows:
        issues.append({"severity": "medium", "code": "duplicate_rows", "column": None, "evidence": f"{duplicate_rows} exact duplicates"})
    completeness = 1 - missing_cells / total_cells
    duplicate_ratio = duplicate_rows / rows
    high_missing_cols = sum(p["missing_pct"] >= 40 for p in column_profiles) / columns
    constant_cols = sum(p["unique"] <= 1 for p in column_profiles) / columns
    score = round(max(0.0, 100 * (1 - (0.55 * (1 - completeness) + 0.2 * duplicate_ratio + 0.15 * high_missing_cols + 0.1 * constant_cols))), 1)
    return {
        "rows": rows, "columns": columns, "cells": total_cells,
        "memory_bytes": int(df.memory_usage(deep=True).sum()),
        "missing_cells": missing_cells, "missing_pct": round((1 - completeness) * 100, 2),
        "duplicate_rows": duplicate_rows, "duplicate_pct": round(duplicate_ratio * 100, 2),
        "quality_score": score, "column_profiles": column_profiles, "issues": issues,
    }


def recommend_cleaning(df: pd.DataFrame, profile: dict[str, Any] | None = None) -> list[dict[str, Any]]:
    """Return conservative, reviewable actions. Destructive choices are never automatic."""
    profile = profile or profile_dataframe(df)
    actions: list[dict[str, Any]] = [
        {"id": "normalize_columns", "action": "normalize_columns", "column": None,
         "label": "Normalize column names", "reason": "Creates stable, SQL-friendly names", "risk": "low", "recommended": True},
        {"id": "normalize_missing", "action": "normalize_missing", "column": None,
         "label": "Normalize missing-value tokens", "reason": "Treats N/A, null, blanks and similar tokens consistently", "risk": "low", "recommended": True},
        {"id": "trim_text", "action": "trim_text", "column": None,
         "label": "Trim surrounding whitespace", "reason": "Prevents duplicate categories caused by spaces", "risk": "low", "recommended": True},
    ]
    if profile["duplicate_rows"]:
        actions.append({"id": "drop_duplicates", "action": "drop_duplicates", "column": None,
                        "label": "Remove exact duplicate rows", "reason": f"Found {profile['duplicate_rows']} exact duplicates", "risk": "medium", "recommended": False})
    for item in profile["column_profiles"]:
        if item["semantic_type"] == "numeric_text":
            actions.append({"id": f"numeric::{item['column']}", "action": "convert_numeric", "column": item["column"],
                            "label": f"Convert {item['column']} to numeric", "reason": f"{item['type_confidence']:.0%} of sampled values parse as numbers", "risk": "medium", "recommended": item["type_confidence"] >= 0.99})
        elif item["semantic_type"] == "datetime_text":
            actions.append({"id": f"datetime::{item['column']}", "action": "convert_datetime", "column": item["column"],
                            "label": f"Convert {item['column']} to datetime", "reason": f"{item['type_confidence']:.0%} of sampled values parse as dates", "risk": "medium", "recommended": item["type_confidence"] >= 0.99})
        if item["unique"] <= 1:
            actions.append({"id": f"drop_constant::{item['column']}", "action": "drop_column", "column": item["column"],
                            "label": f"Drop constant column {item['column']}", "reason": "It contains no analytical variation", "risk": "medium", "recommended": False})
    return actions


def _normalized_column_names(columns: Iterable[Any]) -> list[str]:
    names = []
    for index, column in enumerate(columns, 1):
        name = re.sub(r"[^0-9A-Za-z]+", "_", str(column).strip()).strip("_").lower()
        if not name:
            name = f"column_{index}"
        if name[0].isdigit():
            name = f"col_{name}"
        names.append(name)
    return _dedupe_columns(names)


def apply_cleaning_plan(df: pd.DataFrame, actions: Iterable[dict[str, Any]]) -> tuple[pd.DataFrame, list[dict[str, Any]]]:
    """Apply selected allow-listed actions and return cleaned data plus an audit log."""
    result = df.copy()
    audit: list[dict[str, Any]] = []
    column_map = {str(column): str(column) for column in result.columns}
    allowed = {"normalize_columns", "normalize_missing", "trim_text", "drop_duplicates", "convert_numeric", "convert_datetime", "drop_column"}
    for step, spec in enumerate(actions, 1):
        action = spec.get("action")
        column = spec.get("column")
        if action not in allowed:
            raise ValueError(f"Unsupported cleaning action: {action}")
        before_shape = result.shape
        changed = 0
        if action == "normalize_columns":
            old = list(map(str, result.columns))
            result.columns = _normalized_column_names(result.columns)
            column_map.update(dict(zip(old, result.columns)))
            changed = sum(a != b for a, b in zip(old, result.columns))
        elif action == "normalize_missing":
            object_columns = result.select_dtypes(include=["object", "string"]).columns
            for name in object_columns:
                series = result[name]
                mask = series.notna() & series.astype(str).str.strip().str.lower().isin(MISSING_TOKENS)
                changed += int(mask.sum())
                result.loc[mask, name] = pd.NA
        elif action == "trim_text":
            for name in result.select_dtypes(include=["object", "string"]).columns:
                original = result[name].copy()
                result[name] = result[name].map(lambda value: value.strip() if isinstance(value, str) else value)
                changed += int((original.fillna("<NA>") != result[name].fillna("<NA>")).sum())
        elif action == "drop_duplicates":
            count = len(result)
            result = result.drop_duplicates().reset_index(drop=True)
            changed = count - len(result)
        elif action == "convert_numeric":
            column = column_map.get(str(column), column)
            if column not in result.columns:
                raise ValueError(f"Column no longer exists: {column}")
            parsed = pd.to_numeric(result[column].astype("string").str.replace(",", "", regex=False), errors="coerce")
            changed = int(parsed.notna().sum())
            result[column] = parsed
        elif action == "convert_datetime":
            column = column_map.get(str(column), column)
            if column not in result.columns:
                raise ValueError(f"Column no longer exists: {column}")
            parsed = pd.to_datetime(result[column], errors="coerce", format="mixed")
            changed = int(parsed.notna().sum())
            result[column] = parsed
        elif action == "drop_column":
            column = column_map.get(str(column), column)
            if column not in result.columns:
                raise ValueError(f"Column no longer exists: {column}")
            result = result.drop(columns=[column])
            changed = len(result)
        audit.append({
            "step": step, "action": action, "column": column, "changed": changed,
            "before": list(before_shape), "after": list(result.shape),
            "timestamp": datetime.now(timezone.utc).isoformat(),
        })
    return result, audit


def analyze_dataframe(df: pd.DataFrame, goal: str = "") -> dict[str, Any]:
    """Generate reproducible descriptive findings; never invent causal claims."""
    profile = profile_dataframe(df)
    numeric = df.select_dtypes(include=np.number).columns.tolist()
    categorical = [c for c in df.columns if infer_semantic_type(df[c])[0] == "categorical"]
    findings: list[dict[str, str]] = []
    correlations: list[dict[str, Any]] = []
    if len(numeric) >= 2:
        matrix = df[numeric].corr(method="pearson", min_periods=3)
        for left_index, left in enumerate(numeric):
            for right in numeric[left_index + 1:]:
                value = matrix.loc[left, right]
                if pd.notna(value):
                    correlations.append({"left": str(left), "right": str(right), "correlation": round(float(value), 4), "n": int(df[[left, right]].dropna().shape[0])})
        correlations.sort(key=lambda item: abs(item["correlation"]), reverse=True)
        for item in correlations[:5]:
            if abs(item["correlation"]) >= 0.5:
                findings.append({"type": "relationship", "title": f"{item['left']} and {item['right']} move together",
                                 "evidence": f"Pearson r={item['correlation']:.2f} across {item['n']} complete rows.",
                                 "caveat": "Correlation is descriptive and does not establish causation."})
    for item in profile["column_profiles"]:
        if item.get("outliers_iqr", 0) > 0:
            findings.append({"type": "quality", "title": f"Potential extreme values in {item['column']}",
                             "evidence": f"{item['outliers_iqr']} values fall outside the 1.5×IQR fences.",
                             "caveat": "Extreme values may be valid; investigate before removing them."})
    category_summaries: dict[str, dict[str, int]] = {}
    for column in categorical[:20]:
        counts = df[column].dropna().astype(str).value_counts().head(10)
        category_summaries[str(column)] = {str(k): int(v) for k, v in counts.items()}
        if len(counts):
            findings.append({"type": "distribution", "title": f"Largest {column} group: {counts.index[0]}",
                             "evidence": f"{int(counts.iloc[0])} rows ({counts.iloc[0] / max(df[column].notna().sum(), 1):.1%} of non-missing values).",
                             "caveat": "This is a frequency result, not an outcome comparison."})
    numeric_summary: dict[str, Any] = {}
    for column in numeric:
        series = df[column].dropna()
        numeric_summary[str(column)] = {
            "count": int(series.count()), "mean": _safe_number(series.mean()),
            "median": _safe_number(series.median()), "std": _safe_number(series.std()),
            "min": _safe_number(series.min()), "max": _safe_number(series.max()),
        }
    return {
        "goal": goal.strip(), "profile": profile, "numeric_summary": numeric_summary,
        "categorical_summary": category_summaries, "correlations": correlations[:25],
        "findings": findings[:25],
        "limitations": [
            "Results describe the supplied data and may not generalize beyond it.",
            "Missingness, measurement error, selection bias and confounding can change conclusions.",
            "No causal claim is made without an appropriate research design.",
        ],
    }


def analyze_document_text(text: str) -> dict[str, Any]:
    """Create a transparent baseline analysis for narrative documents.

    This is intentionally extractive rather than generative: every returned
    sentence and number exists in the source text.
    """
    normalized = re.sub(r"[ \t]+", " ", text or "").strip()
    if not normalized:
        return {"words": 0, "characters": 0, "numbers": [], "top_terms": {}, "key_sentences": []}
    words = re.findall(r"\b[A-Za-z][A-Za-z'-]{2,}\b", normalized)
    stopwords = {
        "the", "and", "for", "that", "with", "this", "from", "are", "was", "were",
        "has", "have", "had", "not", "but", "into", "its", "their", "than", "then",
        "page", "can", "will", "would", "could", "should", "about", "also", "they",
    }
    frequencies: dict[str, int] = {}
    for word in words:
        token = word.lower()
        if token not in stopwords:
            frequencies[token] = frequencies.get(token, 0) + 1
    top_terms = dict(sorted(frequencies.items(), key=lambda item: (-item[1], item[0]))[:20])
    number_tokens = re.findall(r"(?<!\w)[+-]?(?:[$£€₦]\s*)?\d[\d,]*(?:\.\d+)?%?", normalized)
    sentences = [segment.strip() for segment in re.split(r"(?<=[.!?])\s+|\n{2,}", normalized) if len(segment.strip()) >= 40]
    scored = []
    for position, sentence in enumerate(sentences[:2_000]):
        tokens = re.findall(r"\b[A-Za-z][A-Za-z'-]{2,}\b", sentence.lower())
        score = sum(frequencies.get(token, 0) for token in tokens) / max(len(tokens), 1)
        scored.append((score, position, sentence))
    selected = sorted(sorted(scored, key=lambda item: item[0], reverse=True)[:8], key=lambda item: item[1])
    return {
        "words": len(words), "characters": len(normalized),
        "numbers": number_tokens[:100], "number_count": len(number_tokens),
        "top_terms": top_terms, "key_sentences": [item[2] for item in selected],
    }


def render_markdown_report(source_name: str, dataset_name: str, analysis: dict[str, Any], audit: list[dict[str, Any]] | None = None) -> str:
    """Create a portable evidence-first report without requiring an LLM."""
    profile = analysis["profile"]
    lines = [
        "# DataPrism Analysis Report", "",
        f"- Source: `{source_name}`", f"- Dataset: `{dataset_name}`",
        f"- Generated: {datetime.now(timezone.utc).isoformat()}",
        f"- Analysis goal: {analysis.get('goal') or 'General exploratory analysis'}", "",
        "## Executive overview", "",
        f"The analyzed table contains **{profile['rows']:,} rows** and **{profile['columns']:,} columns**. "
        f"The deterministic data-quality score is **{profile['quality_score']}/100**, with "
        f"**{profile['missing_pct']:.2f}% missing cells** and **{profile['duplicate_rows']:,} exact duplicate rows**.", "",
        "## Evidence-backed findings", "",
    ]
    if analysis["findings"]:
        for finding in analysis["findings"]:
            lines.extend([f"### {finding['title']}", "", finding["evidence"], "", f"_Caveat: {finding['caveat']}_", ""])
    else:
        lines.extend(["No strong automatic finding met the reporting thresholds.", ""])
    lines.extend(["## Data-quality issues", ""])
    if profile["issues"]:
        for issue in profile["issues"]:
            location = f" in `{issue['column']}`" if issue.get("column") else ""
            lines.append(f"- **{issue['severity'].title()}** — {issue['code'].replace('_', ' ')}{location}: {issue['evidence']}")
    else:
        lines.append("- No rule-based quality issues were detected.")
    lines.extend(["", "## Cleaning provenance", ""])
    if audit:
        for entry in audit:
            location = f" (`{entry['column']}`)" if entry.get("column") else ""
            lines.append(f"- Step {entry['step']}: `{entry['action']}`{location}; changed {entry['changed']} values/rows; shape {tuple(entry['before'])} → {tuple(entry['after'])}.")
    else:
        lines.append("- No cleaning actions were applied.")
    lines.extend(["", "## Limitations", ""] + [f"- {item}" for item in analysis["limitations"]])
    return "\n".join(lines) + "\n"


def export_csv_safe(df: pd.DataFrame) -> bytes:
    """Return UTF-8 CSV while neutralizing spreadsheet formula injection.

    Values beginning with ``=``, ``+``, ``-`` or ``@`` can execute as formulas
    when a downloaded CSV is opened in desktop spreadsheet software. Prefixing
    only those text values with an apostrophe preserves display while preventing
    evaluation. Numeric columns remain numeric.
    """
    safe = df.copy()
    for column in safe.select_dtypes(include=["object", "string"]).columns:
        safe[column] = safe[column].map(
            lambda value: "'" + value if isinstance(value, str) and value.lstrip().startswith(("=", "+", "-", "@")) else value
        )
    return safe.to_csv(index=False).encode("utf-8-sig")
